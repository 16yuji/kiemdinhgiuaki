# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT = Path(r"C:\xampp\htdocs\hotel-booking-travel-mate-leaflet-map-select-result")
ASSET_DIR = PROJECT / "report-assets" / "chapter6"
OUT_PATH = PROJECT / "TravelMate_Chuong6_TrienKhaiChuongTrinh.docx"

FONT = "Times New Roman"
INK = RGBColor(0x1B, 0x1F, 0x24)
NAVY = RGBColor(0x0A, 0x29, 0x38)
MUTED = RGBColor(0x5F, 0x68, 0x73)
BORDER = "D9DDE3"
HEADER_FILL = "EEF3F7"


def set_font(run, size=13, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=13, bold=False, italic=False, color=INK):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
              italic=False, color=INK, before=0, after=6, line=1.3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=NAVY)


def heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=NAVY)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_font(run, size=12, italic=True, color=MUTED)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=13)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_font(run, size=13)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, fill=None, width=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    if width is not None:
        cell.width = width
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    set_font(run, size=12, bold=bold, color=NAVY if bold else INK)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(5.0), Cm(11.0)] if len(rows[0]) == 2 else [Cm(5.0), Cm(5.5), Cm(5.5)]
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            set_cell_text(
                table.cell(i, j),
                text,
                bold=(i == 0),
                fill=HEADER_FILL if i == 0 else None,
                width=widths[j],
            )
    return table


def add_image(doc, path, caption_text, width_in=6.1):
    path = Path(path)
    if not path.exists():
        paragraph(doc, f"[Thiếu ảnh minh họa: {path.name}]", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width_in))
    caption(doc, caption_text)


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    set_style_font(doc.styles["Normal"], size=13)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.3
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[style_name], size=13)
        doc.styles[style_name].paragraph_format.line_spacing = 1.25
        doc.styles[style_name].paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("Báo cáo đồ án cơ sở - Travel Mate")
    set_font(header_run, size=10, italic=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Chương 6 - Triển khai chương trình")
    set_font(footer_run, size=10, italic=True, color=MUTED)

    chapter_title(doc, "CHƯƠNG 6: TRIỂN KHAI CHƯƠNG TRÌNH")
    paragraph(
        doc,
        "Chương này trình bày cách cài đặt, cấu hình và vận hành hệ thống đặt phòng khách sạn trực tuyến Travel Mate. "
        "Nội dung tập trung vào môi trường triển khai, các bước khởi chạy dự án Laravel, dữ liệu mẫu, thanh toán VNPAY sandbox, "
        "bản đồ khách sạn, trợ lý AI cơ bản và minh họa các giao diện chính của hệ thống.",
        after=10,
    )

    heading(doc, "6.1 Hướng dẫn cài đặt và chạy chương trình")
    paragraph(
        doc,
        "Hệ thống Travel Mate được xây dựng theo mô hình website Laravel, sử dụng cơ sở dữ liệu MySQL và giao diện Blade. "
        "Khi triển khai trong môi trường học tập hoặc demo trên máy cá nhân, có thể sử dụng XAMPP để quản lý Apache/MySQL "
        "và chạy ứng dụng bằng lệnh artisan của Laravel.",
    )
    add_table(
        doc,
        [
            ["Thành phần", "Cấu hình sử dụng"],
            ["Ngôn ngữ và framework", "PHP, Laravel, Blade Template, JavaScript, CSS"],
            ["Cơ sở dữ liệu", "MySQL, database demo: hotel_booking"],
            ["Máy chủ local", "XAMPP hoặc PHP built-in server qua php artisan serve"],
            ["Thanh toán", "VNPAY sandbox và thanh toán giả lập phục vụ demo"],
            ["Bản đồ", "Leaflet kết hợp lớp nền OpenStreetMap để hiển thị vị trí khách sạn"],
            ["Tác vụ nền", "Scheduler Laravel xử lý đơn pending_payment hết hạn giữ phòng"],
        ],
    )
    paragraph(doc, "Các bước triển khai cơ bản được thực hiện như sau:", after=4)
    steps = [
        "Cài đặt XAMPP, PHP và Composer phù hợp với phiên bản Laravel của dự án.",
        r"Sao chép thư mục dự án vào thư mục htdocs, ví dụ: C:\xampp\htdocs\hotel-booking-travel-mate-leaflet-map-select-result.",
        "Tạo database MySQL tên hotel_booking trong phpMyAdmin hoặc MySQL client.",
        "Cấu hình file .env theo môi trường local, trong đó APP_URL có thể đặt là http://127.0.0.1:8000 và thông tin kết nối database dùng tài khoản MySQL local.",
        "Chạy composer install để cài dependency backend, sau đó tạo key ứng dụng nếu cần bằng php artisan key:generate.",
        "Khởi tạo lại database bằng php artisan migrate:fresh --seed để có dữ liệu mẫu khách sạn, phòng, tài khoản và giao dịch demo.",
        "Chạy php artisan storage:link để đảm bảo hình ảnh upload/public được truy cập đúng.",
        "Xóa cache cấu hình và view bằng php artisan optimize:clear, php artisan view:clear.",
        "Khởi động website bằng php artisan serve --host=127.0.0.1 --port=8000 rồi truy cập http://127.0.0.1:8000 trên trình duyệt.",
    ]
    for step in steps:
        add_number(doc, step)
    paragraph(
        doc,
        "Đối với thanh toán VNPAY sandbox, hệ thống vẫn giữ luồng thanh toán thật qua cổng sandbox và một luồng thanh toán giả lập để phục vụ demo. "
        "Khi cần test IPN từ VNPAY về máy local, có thể chạy ngrok http 8000 và cấu hình VNPAY_IPN_URL trỏ về đường dẫn /payments/vnpay/ipn trên domain ngrok. "
        "Đường dẫn RETURN_URL có thể dùng http://127.0.0.1:8000/payments/vnpay/return để người dùng quay về đúng phiên đăng nhập trong môi trường local.",
    )
    paragraph(
        doc,
        "Đối với tác vụ giữ phòng 15 phút, khi chạy demo dài hơn nên bật Laravel scheduler bằng php artisan schedule:work để command bookings:expire-pending "
        "được xử lý định kỳ. Nhờ đó các booking pending_payment quá hạn sẽ được giải phóng khỏi kiểm tra tồn phòng.",
    )

    heading(doc, "6.2 Giao diện khách hàng")
    paragraph(
        doc,
        "Giao diện khách hàng được thiết kế theo phong cách Travel Mate với màu nền sáng, điều hướng rõ ràng, ảnh khách sạn lớn và các nút hành động dễ nhận biết. "
        "Khách vãng lai có thể xem trang chủ, danh sách khách sạn và chi tiết khách sạn; khách hàng đăng nhập có thể đặt phòng, thanh toán, theo dõi lịch sử đặt phòng, "
        "gửi yêu cầu hủy/hoàn tiền và đánh giá sau khi hoàn tất lưu trú.",
    )
    add_image(doc, ASSET_DIR / "01-home.png", "Hình 6.1 Giao diện trang chủ Travel Mate", 6.1)
    paragraph(
        doc,
        "Trang chủ đặt trọng tâm vào ô tìm kiếm nhanh theo điểm đến, ngày nhận phòng, ngày trả phòng và số khách. Các thông tin như VNPAY sandbox, "
        "giữ phòng 15 phút và hotline hỗ trợ được đặt ở khu vực dễ quan sát để người dùng hiểu nhanh điều kiện đặt phòng.",
    )
    add_image(doc, ASSET_DIR / "02-hotels.png", "Hình 6.2 Giao diện danh sách khách sạn và bộ lọc", 6.1)
    paragraph(
        doc,
        "Trang danh sách khách sạn hỗ trợ lọc theo điểm đến, thời gian lưu trú, số khách, khoảng giá, đánh giá và tiện nghi. "
        "Kết quả hiển thị theo dạng thẻ khách sạn với ảnh đại diện, vị trí, mức giá, điểm đánh giá và nút xem chi tiết.",
    )
    add_image(doc, ASSET_DIR / "03-hotel-detail.png", "Hình 6.3 Giao diện chi tiết khách sạn", 6.1)
    paragraph(
        doc,
        "Trang chi tiết khách sạn hiển thị tên khách sạn, địa chỉ, đánh giá, thư viện ảnh, tiện nghi, chính sách hủy và khu vực kiểm tra phòng. "
        "Theo quy tắc nghiệp vụ, chỉ tài khoản Customer mới được tạo booking; Admin và Owner có thể xem trang công khai nhưng không được đặt phòng.",
    )
    add_image(doc, ASSET_DIR / "05-map-section.png", "Hình 6.4 Khu vực bản đồ vị trí khách sạn", 5.8)
    paragraph(
        doc,
        "Bản đồ được dùng để giúp khách hàng kiểm tra trực quan vị trí khách sạn. Khi chủ khách sạn nhập hoặc chọn địa chỉ trong phần quản lý, "
        "hệ thống lưu tọa độ và địa chỉ đầy đủ; trang chi tiết hiển thị marker tương ứng trên bản đồ cùng dòng địa chỉ bên dưới.",
    )
    add_image(doc, ASSET_DIR / "06-room-types-section.png", "Hình 6.5 Khu vực hạng phòng và kiểm tra khả dụng", 5.8)
    paragraph(
        doc,
        "Các hạng phòng hiển thị ảnh, tên phòng, mô tả, giá mỗi đêm, sức chứa và các tiện nghi nổi bật. Hệ thống chỉ cho đặt khi còn phòng khả dụng "
        "trong khoảng ngày tìm kiếm; booking pending_payment còn hạn giữ phòng vẫn được tính là đã giữ chỗ để hạn chế overbooking.",
    )
    add_image(doc, ASSET_DIR / "04-login.png", "Hình 6.6 Giao diện đăng nhập người dùng", 5.8)
    paragraph(
        doc,
        "Trang đăng nhập được thiết kế lại theo bố cục hai cột: một bên là hình ảnh thương hiệu, một bên là form đăng nhập. "
        "Hệ thống hỗ trợ đăng nhập bằng email/mật khẩu và nút đăng nhập Google, đồng thời giữ nguyên CSRF token và luồng xác thực sẵn có của Laravel.",
    )

    heading(doc, "6.3 Giao diện chủ khách sạn")
    paragraph(
        doc,
        "Sau khi đăng nhập với vai trò Owner, người dùng được chuyển đến khu vực quản lý dành cho đối tác khách sạn. "
        "Các màn hình trong nhóm này tập trung vào vận hành lưu trú, không cho Owner xử lý hoàn tiền thay Admin.",
    )
    owner_items = [
        "Quản lý khách sạn: tạo, cập nhật thông tin khách sạn, địa chỉ, ảnh, tiện nghi, chính sách hủy và trạng thái kiểm duyệt.",
        "Quản lý hạng phòng và phòng vật lý: khai báo sức chứa, giá phòng, tiện nghi, số lượng phòng và trạng thái phòng.",
        "Quản lý booking: xem danh sách đơn theo trạng thái, xác nhận vận hành, xử lý check-in, check-out, no-show và đổi phòng khi cần.",
        "Check-in chỉ được thực hiện với booking đã xác nhận và vào đúng ngày nhận phòng hoặc sau ngày nhận phòng; nếu không đủ phòng vật lý phù hợp, booking được đưa vào trạng thái manual_review.",
        "Doanh thu: Owner xem doanh thu, giao dịch chờ đối soát, khoản đã thanh toán và các khoản điều chỉnh phát sinh từ hoàn tiền sau khi đã settlement.",
    ]
    for item in owner_items:
        add_bullet(doc, item)

    heading(doc, "6.4 Giao diện quản trị viên")
    paragraph(
        doc,
        "Admin quản lý toàn bộ hệ thống, kiểm soát dữ liệu người dùng, đối tác, khách sạn, hoàn tiền và đối soát. "
        "Các trang quản trị dùng layout thống nhất với sidebar, topbar, bảng dữ liệu, badge trạng thái và form thao tác rõ ràng.",
    )
    admin_items = [
        "Quản lý người dùng: xem danh sách tài khoản, phân loại vai trò Customer, Owner, Admin và theo dõi trạng thái tài khoản.",
        "Duyệt yêu cầu đối tác: tiếp nhận yêu cầu trở thành chủ khách sạn, phê duyệt hoặc từ chối theo thông tin đăng ký.",
        "Kiểm duyệt khách sạn: xem thông tin khách sạn do Owner tạo, duyệt hiển thị công khai hoặc yêu cầu chỉnh sửa.",
        "Quản lý hoàn tiền: xem yêu cầu hủy/hoàn tiền, lý do hủy, chính sách khách sạn, trạng thái thanh toán và nhập số tiền hoàn nếu được duyệt.",
        "Đối soát Owner: xác nhận settlement, hiển thị tổng doanh thu Owner, các khoản khấu trừ pending adjustment và số tiền thực chuyển.",
        "Kiểm duyệt đánh giá: quản lý review của khách hàng và phản hồi của Owner để đảm bảo nội dung phù hợp.",
    ]
    for item in admin_items:
        add_bullet(doc, item)

    doc.add_page_break()
    heading(doc, "6.5 Kiểm thử các luồng nghiệp vụ chính")
    paragraph(
        doc,
        "Trong quá trình triển khai, hệ thống được kiểm tra theo các luồng nghiệp vụ trọng tâm để bảo đảm giao diện không làm sai logic backend. "
        "Các luồng quan trọng gồm đặt phòng, giữ phòng, thanh toán, hủy đơn, hoàn tiền và đối soát.",
    )
    add_table(
        doc,
        [
            ["Luồng kiểm thử", "Kết quả mong đợi"],
            ["Giữ phòng 15 phút", "Booking pending_payment còn hạn hold_expires_at được tính là đã giữ chỗ; booking quá hạn không còn giữ phòng."],
            ["Thanh toán VNPAY sandbox", "Khi callback hợp lệ và đúng số tiền, booking chuyển confirmed, payment chuyển paid và giao dịch tài chính chỉ được tạo một lần."],
            ["Thanh toán giả lập", "Khách hàng có thể dùng fallback demo để chuyển booking sang trạng thái đã thanh toán trong môi trường test."],
            ["Hủy booking chưa thanh toán", "Customer có thể hủy trước ngày check-in; booking chuyển cancelled, payment pending chuyển failed/expired và không phát sinh hoàn tiền."],
            ["Hủy booking đã thanh toán", "Customer gửi yêu cầu hủy/hoàn tiền trước ngày check-in; Admin xem xét theo chính sách khách sạn."],
            ["Hoàn tiền sau settlement", "Nếu giao dịch đã settlement, hệ thống tạo OwnerAdjustment dạng clawback và trừ vào settlement tiếp theo của cùng Owner."],
            ["Check-in/check-out", "Owner chỉ check-in booking confirmed đúng ngày hoặc sau ngày nhận phòng; check-out chỉ áp dụng cho booking staying."],
            ["Bảo vệ dữ liệu tài chính", "Trang Customer chỉ hiển thị tổng tiền booking, trạng thái thanh toán/hoàn tiền và thông tin hỗ trợ, không hiển thị dữ liệu nội bộ."],
        ],
    )
    paragraph(
        doc,
        "Bên cạnh các luồng đặt phòng và thanh toán, hệ thống đã bổ sung bản đồ cơ bản và trợ lý AI hỗ trợ người dùng. "
        "Bản đồ phục vụ việc hiển thị vị trí khách sạn và hỗ trợ chọn địa chỉ trong phần quản lý. Trợ lý AI dùng dữ liệu hiện có của hệ thống để trả lời "
        "các câu hỏi như gợi ý khách sạn, số hạng phòng, khoảng giá, tiện nghi và chính sách hủy của khách sạn đang xem. "
        "Đây là mức tích hợp phục vụ demo và có thể tiếp tục mở rộng ở các phiên bản sau.",
    )

    heading(doc, "6.6 Nhận xét triển khai")
    paragraph(
        doc,
        "Sau khi hoàn thiện, Travel Mate đáp ứng được các vai trò chính gồm Guest, Customer, Owner và Admin. Người dùng có thể tìm khách sạn, xem chi tiết, "
        "đặt phòng, thanh toán sandbox/demo, theo dõi booking và liên hệ hỗ trợ khi cần. Owner có công cụ vận hành khách sạn, còn Admin có công cụ kiểm duyệt, "
        "hoàn tiền và đối soát.",
    )
    paragraph(
        doc,
        "Phần triển khai hiện phù hợp cho môi trường học tập, demo và kiểm thử nghiệp vụ. Khi đưa lên môi trường production, cần cấu hình domain thật, SSL, "
        "queue/scheduler chạy nền ổn định, tài khoản VNPAY chính thức, chính sách giới hạn request cho geocoding/bản đồ và giám sát log thanh toán để bảo đảm độ tin cậy lâu dài.",
    )

    doc.core_properties.title = "Chương 6 - Triển khai chương trình Travel Mate"
    doc.core_properties.subject = "Báo cáo đồ án cơ sở"
    doc.core_properties.author = "Travel Mate"
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
