# -*- coding: utf-8 -*-
import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT = Path(r"C:\xampp\htdocs\hotel-booking-travel-mate-leaflet-map-select-result")
SCREEN_DIR = PROJECT / "report-assets" / "full-ui-screens"
OPTIMIZED_DIR = PROJECT / "report-assets" / "chapter5-combined-screens-docx"
MAIN_MANIFEST = SCREEN_DIR / "screens_manifest.json"
AI_MANIFEST = SCREEN_DIR / "ai_chatbot_manifest.json"
OUT_PATH = Path(r"E:\abc\TravelMate_Chuong5_TestCase_GiaoDien.docx")

FONT = "Times New Roman"
INK = RGBColor(0x1B, 0x1F, 0x24)
NAVY = RGBColor(0x0A, 0x29, 0x38)
MUTED = RGBColor(0x5F, 0x68, 0x73)
GREEN = RGBColor(0x16, 0x6A, 0x3B)
BORDER = "D9DDE3"
HEADER_FILL = "EEF3F7"
PASS_FILL = "EAF7EF"

ROLE_TITLES = {
    "public": "5.3.1 Giao diện công khai và xác thực",
    "customer": "5.3.2 Giao diện khách hàng",
    "owner": "5.3.3 Giao diện chủ khách sạn",
    "admin": "5.3.4 Giao diện quản trị viên",
    "ai": "5.3.5 Giao diện chatbot Travel Mate AI",
}

ROLE_DESCRIPTIONS = {
    "public": "Nhóm giao diện dành cho khách vãng lai: trang chủ, tìm khách sạn, chi tiết khách sạn và các màn hình xác thực.",
    "customer": "Nhóm giao diện dành cho khách hàng: tìm phòng, đặt phòng, thanh toán, lịch sử đặt phòng, hồ sơ và yêu cầu đối tác.",
    "owner": "Nhóm giao diện dành cho chủ khách sạn: quản lý khách sạn, hạng phòng, phòng vật lý, đơn đặt phòng, vận hành lưu trú, doanh thu và đánh giá.",
    "admin": "Nhóm giao diện dành cho quản trị viên: quản lý người dùng, đối tác, khách sạn, hoàn tiền, đối soát, đánh giá và khiếu nại.",
    "ai": "Giao diện chatbot hỗ trợ tư vấn nhanh về khách sạn, đặt phòng, thanh toán, hủy/hoàn tiền và thao tác trên Travel Mate.",
}

TEST_CASES = [
    ["Ca 01", "Cao", "Đăng ký và đăng nhập", "Người dùng nhập email, mật khẩu và thông tin hợp lệ.", "Tài khoản được tạo hoặc đăng nhập thành công, hệ thống đưa người dùng đến đúng khu vực theo vai trò."],
    ["Ca 02", "Cao", "Phân quyền vai trò", "Người dùng thử truy cập các khu vực không thuộc quyền của mình.", "Hệ thống chặn đúng quyền; quản trị viên và chủ khách sạn có thể xem trang công khai nhưng không thể đặt phòng như khách hàng."],
    ["Ca 03", "Cao", "Tìm kiếm khách sạn", "Người dùng nhập điểm đến, ngày nhận phòng, ngày trả phòng và số khách.", "Danh sách khách sạn hiển thị đúng điều kiện tìm kiếm, có ảnh, vị trí, đánh giá và giá tham khảo."],
    ["Ca 04", "Cao", "Xem chi tiết khách sạn", "Người dùng mở một khách sạn đang hoạt động, có ảnh, hạng phòng, chính sách và vị trí bản đồ.", "Trang chi tiết hiển thị đầy đủ bộ ảnh, tiện nghi, chính sách hủy, bản đồ và các hạng phòng còn phù hợp."],
    ["Ca 05", "Cao", "Tạo đơn đặt phòng", "Khách hàng chọn hạng phòng còn trống, ngày hợp lệ và nhập thông tin liên hệ.", "Đơn đặt phòng được tạo ở trạng thái chờ thanh toán, hệ thống giữ phòng tạm thời trong 15 phút."],
    ["Ca 06", "Cao", "Chống đặt phòng vượt số lượng", "Hai khách hàng cùng đặt số phòng cuối cùng của một hạng phòng.", "Hệ thống kiểm tra lại số phòng trước khi xác nhận, không cho tạo số đơn vượt quá số phòng thật."],
    ["Ca 07", "Cao", "Hết hạn giữ phòng", "Đơn đặt phòng chưa thanh toán đã quá thời gian giữ phòng.", "Hệ thống tự chuyển đơn sang trạng thái hết hạn, giao dịch chưa thanh toán không còn hiệu lực và phòng được mở lại."],
    ["Ca 08", "Cao", "Thanh toán", "Khách hàng thanh toán bằng phương thức giả lập hoặc qua VNPAY sandbox.", "Thanh toán thành công thì đơn được xác nhận, trạng thái thanh toán được cập nhật và hệ thống không ghi nhận trùng khi nhận thông báo nhiều lần."],
    ["Ca 09", "Cao", "Hủy đơn chưa thanh toán", "Khách hàng có đơn chưa thanh toán và ngày nhận phòng vẫn ở tương lai.", "Khách hàng có thể hủy ngay; đơn chuyển sang đã hủy, không phát sinh hoàn tiền."],
    ["Ca 10", "Cao", "Hủy đơn đã thanh toán", "Khách hàng có đơn đã thanh toán và ngày nhận phòng vẫn ở tương lai.", "Khách hàng gửi yêu cầu hủy và hoàn tiền; đơn chuyển sang đã hủy và chờ quản trị viên xem xét hoàn tiền."],
    ["Ca 11", "Cao", "Chặn hủy trực tuyến vào ngày nhận phòng", "Đơn đã thanh toán có ngày nhận phòng là hôm nay hoặc đã qua.", "Trang không hiển thị form hủy trực tuyến; khách hàng được hướng dẫn liên hệ bộ phận hỗ trợ và khách sạn."],
    ["Ca 12", "Trung bình", "Đánh giá sau lưu trú", "Khách hàng có đơn đã hoàn tất và chưa từng đánh giá.", "Khách hàng gửi được một đánh giá; các đơn chưa hoàn tất hoặc đã đánh giá không được gửi lại."],
    ["Ca 13", "Cao", "Chủ khách sạn vận hành lưu trú", "Chủ khách sạn thực hiện nhận phòng, trả phòng hoặc đánh dấu khách không đến.", "Nhận phòng chỉ thực hiện khi đúng ngày và đủ phòng; trả phòng hoàn tất lưu trú; khách không đến chỉ được ghi nhận sau ngày nhận phòng."],
    ["Ca 14", "Cao", "Quản trị viên xử lý hoàn tiền", "Có yêu cầu hoàn tiền từ khách hàng, giao dịch có thể đã hoặc chưa đối soát cho chủ khách sạn.", "Nếu chưa đối soát, hệ thống cập nhật hoàn tiền trực tiếp. Nếu đã đối soát, hệ thống tạo khoản điều chỉnh để khấu trừ ở kỳ sau."],
    ["Ca 15", "Cao", "Đối soát doanh thu", "Chủ khách sạn có doanh thu chờ chuyển và có thể có khoản điều chỉnh.", "Số tiền chuyển thực tế được tính sau khi trừ các khoản điều chỉnh; nếu khoản trừ lớn hơn doanh thu kỳ này thì phần còn lại được giữ sang kỳ sau."],
    ["Ca 16", "Cao", "Bảo mật dữ liệu tài chính", "Khách hàng mở trang thanh toán, lịch sử đặt phòng và chi tiết đơn.", "Khách hàng chỉ thấy tổng tiền, trạng thái thanh toán và thông tin hoàn tiền cần thiết; dữ liệu chia doanh thu nội bộ không hiển thị."],
    ["Ca 17", "Trung bình", "Bản đồ khách sạn", "Khách sạn có địa chỉ chi tiết và vị trí trên bản đồ.", "Bản đồ hiển thị đúng vị trí khách sạn, địa chỉ rõ ràng; nếu bản đồ tải chậm thì luồng đặt phòng vẫn không bị ảnh hưởng."],
    ["Ca 18", "Trung bình", "Chatbot AI", "Người dùng hỏi về khách sạn, giá, hạng phòng hoặc chính sách hủy.", "Chatbot mở đúng giao diện, trả lời theo ngữ cảnh trang và hướng dẫn người dùng thao tác tiếp theo."],
]


def set_font(run, size=13, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=13, bold=False, italic=False, color=INK):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
              italic=False, color=INK, before=0, after=6, line=1.3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=NAVY)


def heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 2 else 7)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=14 if level == 2 else 13, bold=True, color=NAVY)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=12)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=12, italic=True, color=MUTED)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, fill=None, size=10.5, color=INK,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    for run in p.runs:
        run.text = ""
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths_cm, font_size=10.5):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    set_table_width(table, total_dxa)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, fill=HEADER_FILL, size=font_size, color=NAVY,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, value in enumerate(row_data):
            fill = PASS_FILL if headers[c_idx] == "Đánh giá" else None
            color = GREEN if headers[c_idx] == "Đánh giá" else INK
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 1, len(headers) - 1] else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(r_idx, c_idx), value, size=font_size, fill=fill, color=color,
                          align=align, bold=(headers[c_idx] == "Đánh giá"))
    paragraph(doc, "", after=4)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    set_style_font(doc.styles["Normal"], size=13)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.3
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[style_name], size=12)
        doc.styles[style_name].paragraph_format.line_spacing = 1.2
        doc.styles[style_name].paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(header.add_run("Báo cáo đồ án cơ sở - Travel Mate"), size=10, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Chương 5 - Cài đặt thực nghiệm, giao diện và kiểm thử hệ thống"), size=10, italic=True, color=MUTED)
    return doc


def optimize_image(src):
    OPTIMIZED_DIR.mkdir(parents=True, exist_ok=True)
    src = Path(src)
    dest = OPTIMIZED_DIR / src.name
    if dest.exists() and dest.stat().st_mtime >= src.stat().st_mtime:
        return dest
    with Image.open(src) as image:
        image = image.convert("RGB")
        max_width = 1180
        if image.width > max_width:
            ratio = max_width / image.width
            image = image.resize((max_width, int(image.height * ratio)), Image.LANCZOS)
        image.save(dest, format="PNG", optimize=True)
    return dest


def add_screen(doc, item, figure_number):
    heading(doc, item["title"], level=3)
    image_path = optimize_image(item["file"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(6.25))
    caption(doc, f"Hình 5.{figure_number} {item['caption']}")


def load_screens():
    data = []
    if MAIN_MANIFEST.exists():
        data.extend(json.loads(MAIN_MANIFEST.read_text(encoding="utf-8")))
    if AI_MANIFEST.exists():
        data.extend(json.loads(AI_MANIFEST.read_text(encoding="utf-8")))

    clean = []
    seen = set()
    for item in data:
        file_name = item.get("fileName")
        if not file_name or file_name in seen:
            continue
        if "error" in item:
            continue
        if file_name == "customer-09-payment-info.png":
            continue
        if not Path(item.get("file", "")).exists():
            continue
        seen.add(file_name)
        clean.append(item)
    return clean


def build_document():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    screens = load_screens()
    grouped = {}
    for item in screens:
        grouped.setdefault(item["role"], []).append(item)

    doc = setup_doc()
    chapter_title(doc, "CHƯƠNG 5: CÀI ĐẶT THỰC NGHIỆM, GIAO DIỆN VÀ KIỂM THỬ HỆ THỐNG")
    paragraph(
        doc,
        "Chương này trình bày quá trình cài đặt thực nghiệm, minh họa giao diện và kết quả kiểm thử sơ bộ của hệ thống Travel Mate. "
        "Phần đầu tập trung vào môi trường triển khai và các màn hình chính của hệ thống; phần kiểm thử được đặt ở cuối chương để đối chiếu lại các chức năng trọng tâm sau khi đã quan sát giao diện thực tế.",
    )

    heading(doc, "5.1 Cài đặt thực nghiệm")
    paragraph(
        doc,
        "Hệ thống được chạy thử trên môi trường máy cá nhân với Laravel, PHP, MySQL/MariaDB và dữ liệu mẫu đã được chuẩn bị sẵn. "
        "Mục tiêu của phần cài đặt thực nghiệm là tạo đủ dữ liệu cho các vai trò Guest/Public, Customer, Owner và Admin, đồng thời kiểm tra được các luồng đặt phòng, thanh toán, vận hành khách sạn và quản trị hệ thống.",
    )
    add_table(
        doc,
        ["Thành phần", "Cấu hình thực nghiệm"],
        [
            ["Dự án", "Travel Mate được triển khai dưới dạng website Laravel phục vụ đặt phòng khách sạn trực tuyến."],
            ["Máy chủ thử nghiệm", "Hệ thống được chạy trên máy cá nhân để phục vụ demo và kiểm thử chức năng."],
            ["Cơ sở dữ liệu", "Dữ liệu được lưu trong MySQL/MariaDB thông qua môi trường XAMPP."],
            ["Dữ liệu mẫu", "Dữ liệu mẫu bao gồm khách sạn, ảnh, hạng phòng, phòng vật lý, đơn đặt phòng, thanh toán, đối soát và khoản điều chỉnh doanh thu."],
            ["Thanh toán", "Hệ thống giữ cả thanh toán VNPAY sandbox và thanh toán giả lập để thuận tiện cho demo."],
            ["Bản đồ và AI", "Bản đồ dùng thư viện hiển thị vị trí khách sạn; chatbot Travel Mate AI hiển thị trên giao diện để tư vấn theo ngữ cảnh."],
        ],
        [4.8, 12.2],
        font_size=10.5,
    )
    add_bullet(doc, "Chuẩn bị các thư viện cần thiết cho Laravel và giao diện trước khi chạy thử hệ thống.")
    add_bullet(doc, "Cấu hình thông tin môi trường như cơ sở dữ liệu, email, thanh toán và địa chỉ website thử nghiệm.")
    add_bullet(doc, "Tạo lại cấu trúc dữ liệu và nạp dữ liệu mẫu để có đủ tài khoản, khách sạn, phòng và đơn đặt phòng phục vụ kiểm thử.")
    add_bullet(doc, "Làm mới bộ nhớ đệm của ứng dụng rồi khởi động website trên máy cá nhân.")
    add_bullet(doc, "Khi cần kiểm thử thông báo thanh toán từ VNPAY về máy cá nhân, có thể tạo một địa chỉ truy cập tạm thời từ bên ngoài để hệ thống thanh toán gửi kết quả về.")

    heading(doc, "5.2 Phạm vi thực nghiệm")
    add_table(
        doc,
        ["Nội dung", "Mô tả"],
        [
            ["Vai trò", "Guest/Public, Customer, Owner và Admin."],
            ["Public/Auth", "Trang chủ, tìm khách sạn, chi tiết khách sạn, đăng nhập, đăng ký, Google OAuth và quên mật khẩu."],
            ["Customer", "Tìm phòng, đặt phòng, thanh toán, lịch sử đặt phòng, hủy/hoàn tiền, đánh giá, yêu cầu đối tác và hồ sơ cá nhân."],
            ["Owner", "Quản lý khách sạn, hạng phòng, phòng vật lý, đơn đặt phòng, nhận phòng, trả phòng, đánh dấu khách không đến, đổi phòng, doanh thu và đánh giá."],
            ["Admin", "Quản lý người dùng, đối tác, kiểm duyệt khách sạn, hoàn tiền, đối soát, đánh giá và khiếu nại trạng thái khách sạn."],
            ["Tích hợp", "VNPAY sandbox, thanh toán giả lập demo, bản đồ hiển thị vị trí khách sạn và chatbot Travel Mate AI."],
            ["Nguyên tắc bảo mật", "Khách hàng chỉ xem được thông tin thanh toán cần thiết, không thấy dữ liệu chia doanh thu nội bộ của nền tảng và chủ khách sạn."],
        ],
        [4.2, 12.8],
        font_size=10.5,
    )

    doc.add_page_break()
    heading(doc, "5.3 Minh họa giao diện hệ thống")
    paragraph(
        doc,
        "Phần này minh họa các giao diện chính sau khi hệ thống được cài đặt thực nghiệm. "
        "Các ảnh được chia theo nhóm vai trò và đã bổ sung ảnh chatbot Travel Mate AI để thể hiện phần tư vấn người dùng.",
    )

    figure = 1
    for role in ["public", "customer", "owner", "admin", "ai"]:
        items = grouped.get(role, [])
        if not items:
            continue
        heading(doc, ROLE_TITLES[role])
        paragraph(doc, ROLE_DESCRIPTIONS[role])
        for item in items:
            add_screen(doc, item, figure)
            figure += 1
        if role != "ai":
            doc.add_page_break()

    doc.add_page_break()
    heading(doc, "5.4 Kiểm thử chức năng chính")
    paragraph(
        doc,
        "Sau khi cài đặt và quan sát giao diện, hệ thống được kiểm thử sơ bộ theo các chức năng chính. "
        "Danh sách dưới đây tập trung vào các luồng quan trọng, không đi quá sâu vào từng biến thể nhỏ để giữ chương báo cáo gọn và dễ theo dõi.",
    )
    heading(doc, "5.4.1 Nhóm Public/Auth và Customer", level=3)
    test_rows = [row + ["Đạt"] for row in TEST_CASES[:12]]
    add_table(
        doc,
        ["Ca kiểm thử", "Mức", "Chức năng", "Dữ liệu / điều kiện", "Kết quả mong đợi", "Đánh giá"],
        test_rows,
        [1.8, 1.5, 3.4, 5.1, 5.1, 1.1],
        font_size=9.2,
    )
    heading(doc, "5.4.2 Nhóm Owner, Admin và tích hợp", level=3)
    test_rows = [row + ["Đạt"] for row in TEST_CASES[12:]]
    add_table(
        doc,
        ["Ca kiểm thử", "Mức", "Chức năng", "Dữ liệu / điều kiện", "Kết quả mong đợi", "Đánh giá"],
        test_rows,
        [1.8, 1.5, 3.4, 5.1, 5.1, 1.1],
        font_size=9.2,
    )

    heading(doc, "5.5 Tổng hợp kết quả kiểm thử")
    add_table(
        doc,
        ["Nhóm chức năng", "Số test case", "Kết quả"],
        [
            ["Public/Auth và phân quyền", "4", "Đạt"],
            ["Khách hàng đặt phòng, thanh toán, hủy và đánh giá", "8", "Đạt"],
            ["Owner vận hành khách sạn", "2", "Đạt"],
            ["Quản trị viên hoàn tiền, đối soát và kiểm duyệt", "3", "Đạt"],
            ["Map và chatbot AI", "2", "Đạt"],
            ["Tổng cộng", str(len(TEST_CASES)), "Đạt"],
        ],
        [7.2, 3.0, 6.6],
        font_size=10.5,
    )

    heading(doc, "5.6 Nhận xét")
    paragraph(
        doc,
        "Bố cục chương đã đặt phần cài đặt thực nghiệm và giao diện lên trước để người đọc thấy được cách hệ thống được triển khai và vận hành thực tế. "
        "Phần kiểm thử đặt ở cuối chương giúp đối chiếu lại các chức năng cốt lõi sau khi đã xem giao diện. Kết quả kiểm thử sơ bộ cho thấy Travel Mate đáp ứng các luồng chính về tìm khách sạn, đặt phòng, thanh toán, hủy/hoàn tiền, vận hành khách sạn, hoàn tiền sau đối soát, bảo mật dữ liệu tài chính và hỗ trợ chatbot AI.",
    )

    doc.core_properties.title = "Chương 5 - Cài đặt thực nghiệm, giao diện và kiểm thử Travel Mate"
    doc.core_properties.subject = "Báo cáo đồ án cơ sở"
    doc.core_properties.author = "Travel Mate"
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
