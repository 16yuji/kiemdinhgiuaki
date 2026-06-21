# -*- coding: utf-8 -*-
import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT = Path(r"C:\xampp\htdocs\hotel-booking-travel-mate-leaflet-map-select-result")
SCREEN_DIR = PROJECT / "report-assets" / "full-ui-screens"
OPTIMIZED_DIR = PROJECT / "report-assets" / "full-ui-screens-docx"
MANIFEST_PATH = SCREEN_DIR / "screens_manifest.json"
OUT_PATH = PROJECT / "TravelMate_Chuong6_ToanBoGiaoDien.docx"

FONT = "Times New Roman"
INK = RGBColor(0x1B, 0x1F, 0x24)
NAVY = RGBColor(0x0A, 0x29, 0x38)
MUTED = RGBColor(0x5F, 0x68, 0x73)
BORDER = "D9DDE3"
HEADER_FILL = "EEF3F7"

ROLE_TITLES = {
    "public": "6.2 Giao diện công khai và xác thực",
    "customer": "6.3 Giao diện khách hàng",
    "owner": "6.4 Giao diện chủ khách sạn",
    "admin": "6.5 Giao diện quản trị viên",
}

ROLE_DESCRIPTIONS = {
    "public": (
        "Nhóm giao diện này dành cho khách vãng lai và các màn hình xác thực. Người dùng có thể xem landing page, "
        "tìm kiếm khách sạn, xem chi tiết khách sạn, đăng nhập, đăng ký và khôi phục mật khẩu."
    ),
    "customer": (
        "Nhóm giao diện khách hàng phục vụ luồng tìm phòng, xem chi tiết khách sạn, đặt phòng, thanh toán, theo dõi lịch sử, "
        "yêu cầu trở thành đối tác và quản lý hồ sơ cá nhân."
    ),
    "owner": (
        "Nhóm giao diện Owner hỗ trợ chủ khách sạn vận hành cơ sở lưu trú, bao gồm quản lý khách sạn, hạng phòng, phòng vật lý, "
        "booking, check-in, đổi phòng, doanh thu và phản hồi đánh giá."
    ),
    "admin": (
        "Nhóm giao diện Admin hỗ trợ quản trị toàn hệ thống, gồm người dùng, đối tác, kiểm duyệt khách sạn, hoàn tiền, đối soát, "
        "đánh giá và yêu cầu xem xét trạng thái khách sạn."
    ),
}


def set_font(run, size=13, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=13, bold=False, italic=False, color=INK):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
              italic=False, color=INK, before=0, after=6, line=1.3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=NAVY)


def heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=14 if level == 2 else 13, bold=True, color=NAVY)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=12, italic=True, color=MUTED)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=12)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    set_font(run, size=11, bold=bold, color=NAVY if bold else INK)


def add_summary_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(5.2), Cm(6.2), Cm(5.4)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
    headers = ["Nhóm giao diện", "Số màn hình", "Mục đích"]
    for j, header in enumerate(headers):
        set_cell_text(table.cell(0, j), header, bold=True, fill=HEADER_FILL)
    for i, row in enumerate(rows, start=1):
        for j, text in enumerate(row):
            set_cell_text(table.cell(i, j), text)


def optimize_image(src):
    OPTIMIZED_DIR.mkdir(parents=True, exist_ok=True)
    src = Path(src)
    dest = OPTIMIZED_DIR / src.name
    if dest.exists() and dest.stat().st_mtime >= src.stat().st_mtime:
        return dest
    with Image.open(src) as image:
        image = image.convert("RGB")
        max_width = 1180
        if image.width > max_width:
            ratio = max_width / image.width
            image = image.resize((max_width, int(image.height * ratio)), Image.LANCZOS)
        image.save(dest, format="PNG", optimize=True)
    return dest


def add_screen(doc, item, figure_number):
    image_path = Path(item["file"])
    optimized = optimize_image(image_path)
    heading(doc, f'{item["title"]}', level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(optimized), width=Inches(6.25))
    caption(doc, f'Hình 6.{figure_number} {item["caption"]}')


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    set_style_font(doc.styles["Normal"], size=13)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.3
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[style_name], size=12)
        doc.styles[style_name].paragraph_format.line_spacing = 1.2
        doc.styles[style_name].paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(header.add_run("Báo cáo đồ án cơ sở - Travel Mate"), size=10, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Chương 6 - Toàn bộ giao diện hệ thống"), size=10, italic=True, color=MUTED)
    return doc


def build_document():
    data = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    data = [
        item
        for item in data
        if "error" not in item
        and Path(item["file"]).exists()
        and item.get("fileName") != "customer-09-payment-info.png"
    ]

    grouped = {}
    for item in data:
        grouped.setdefault(item["role"], []).append(item)

    doc = setup_doc()
    chapter_title(doc, "CHƯƠNG 6: TRIỂN KHAI CHƯƠNG TRÌNH VÀ GIAO DIỆN HỆ THỐNG")
    paragraph(
        doc,
        "Chương này trình bày cách triển khai chương trình Travel Mate và tổng hợp toàn bộ các giao diện chính đang có trong dự án. "
        "Các hình minh họa được chụp trực tiếp từ hệ thống Laravel đang chạy ở môi trường local, chia theo nhóm người dùng Guest/Public, Customer, Owner và Admin.",
    )

    heading(doc, "6.1 Hướng dẫn cài đặt và chạy chương trình")
    paragraph(
        doc,
        "Để chạy hệ thống Travel Mate trên máy cá nhân, cần cài đặt XAMPP hoặc môi trường PHP/MySQL tương đương, Composer và cấu hình file .env cho database hotel_booking. "
        "Sau đó chạy composer install, php artisan key:generate nếu cần, php artisan migrate:fresh --seed để tạo dữ liệu mẫu, php artisan storage:link để liên kết thư mục lưu ảnh, "
        "php artisan optimize:clear để xóa cache và php artisan serve --host=127.0.0.1 --port=8000 để mở website tại http://127.0.0.1:8000."
    )
    paragraph(
        doc,
        "Với thanh toán VNPAY sandbox, hệ thống giữ cả luồng thanh toán qua VNPAY và thanh toán giả lập phục vụ demo. Khi cần kiểm thử IPN từ VNPAY về máy local, có thể dùng ngrok "
        "trỏ về cổng 8000 và cấu hình VNPAY_IPN_URL theo domain ngrok. Tác vụ hết hạn giữ phòng được xử lý qua Laravel scheduler bằng lệnh php artisan schedule:work.",
    )

    add_summary_table(
        doc,
        [
            ["Public/Auth", str(len(grouped.get("public", []))), "Trang công khai, tìm khách sạn và xác thực tài khoản."],
            ["Customer", str(len(grouped.get("customer", []))), "Đặt phòng, thanh toán, lịch sử booking, hồ sơ và yêu cầu đối tác."],
            ["Owner", str(len(grouped.get("owner", []))), "Quản lý khách sạn, phòng, booking, check-in, đổi phòng, doanh thu và đánh giá."],
            ["Admin", str(len(grouped.get("admin", []))), "Quản trị người dùng, đối tác, khách sạn, refund, settlement và review."],
        ],
    )

    figure = 1
    for role in ["public", "customer", "owner", "admin"]:
        doc.add_page_break()
        heading(doc, ROLE_TITLES[role])
        paragraph(doc, ROLE_DESCRIPTIONS[role])
        for item in grouped.get(role, []):
            add_screen(doc, item, figure)
            figure += 1

    doc.add_page_break()
    heading(doc, "6.6 Nhận xét tổng quan giao diện")
    paragraph(
        doc,
        "Các giao diện trong dự án đã được đồng bộ theo phong cách Travel Mate: nền sáng, navy, vàng ấm, bố cục card/bảng rõ ràng, badge trạng thái dễ đọc và điều hướng theo vai trò. "
        "Những màn hình khách hàng không hiển thị dữ liệu tài chính nội bộ như platform_fee, owner_amount hoặc financialTransaction; các thông tin đối soát và khấu trừ chỉ nằm trong khu vực Owner/Admin phù hợp nghiệp vụ.",
    )
    add_bullet(doc, "Guest/Public có thể xem trang chủ, danh sách khách sạn, chi tiết khách sạn và các màn hình xác thực.")
    add_bullet(doc, "Customer có đầy đủ màn hình tìm phòng, đặt phòng, thanh toán VNPAY/demo, xem booking, hủy/hoàn tiền theo điều kiện, hồ sơ và yêu cầu đối tác.")
    add_bullet(doc, "Owner có các màn hình phục vụ vận hành khách sạn: quản lý khách sạn, hạng phòng, phòng vật lý, booking, check-in, đổi phòng, doanh thu và review.")
    add_bullet(doc, "Admin có các màn hình quản trị hệ thống: người dùng, đối tác, khách sạn, refund, settlement, review và yêu cầu xem xét trạng thái.")
    paragraph(
        doc,
        "Bộ ảnh trong chương này có thể dùng trực tiếp làm phần minh họa triển khai chương trình trong báo cáo đồ án. Nếu cần rút gọn khi nộp bản chính, có thể giữ lại các màn hình tiêu biểu theo từng nhóm vai trò và chuyển phần còn lại sang phụ lục.",
    )

    doc.core_properties.title = "Chương 6 - Toàn bộ giao diện Travel Mate"
    doc.core_properties.subject = "Báo cáo đồ án cơ sở"
    doc.core_properties.author = "Travel Mate"
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
