# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT = Path(r"C:\xampp\htdocs\hotel-booking-travel-mate-leaflet-map-select-result")
OUT_PATH = PROJECT / "TravelMate_Chuong5_TestCase.docx"

FONT = "Times New Roman"
INK = RGBColor(0x1B, 0x1F, 0x24)
NAVY = RGBColor(0x0A, 0x29, 0x38)
MUTED = RGBColor(0x5F, 0x68, 0x73)
GREEN = RGBColor(0x16, 0x6A, 0x3B)
RED = RGBColor(0x9B, 0x1C, 0x1C)
BORDER = "D9DDE3"
HEADER_FILL = "EEF3F7"
PASS_FILL = "EAF7EF"
WARN_FILL = "FFF4D6"


def set_font(run, size=13, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size=13, bold=False, italic=False, color=INK):
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=13, bold=False,
              italic=False, color=INK, before=0, after=6, line=1.3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def chapter_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=NAVY)


def heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 2 else 7)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=14 if level == 2 else 13, bold=True, color=NAVY)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size=12)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, fill=None, size=10.5, color=INK,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    if p.runs:
        p.runs[0].text = ""
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def add_simple_table(doc, headers, rows, widths_cm, font_size=10.5):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_dxa = int(sum(widths_cm) / 2.54 * 1440)
    set_table_width(table, total_dxa)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, fill=HEADER_FILL, size=font_size, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, value in enumerate(row_data):
            set_cell_text(table.cell(r_idx, c_idx), str(value), size=font_size)
    return table


def add_test_table(doc, title, rows):
    heading(doc, title, level=3)
    headers = ["Mã TC", "Mức", "Chức năng", "Dữ liệu / điều kiện", "Bước kiểm thử", "Kết quả mong đợi", "KQ"]
    widths = [1.7, 1.1, 3.3, 4.4, 5.3, 5.3, 1.0]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 12480)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths[idx])
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, fill=HEADER_FILL, size=9.2, color=NAVY,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows, start=1):
        values = [row["id"], row["priority"], row["feature"], row["data"], row["steps"], row["expected"], row["result"]]
        for c_idx, value in enumerate(values):
            fill = PASS_FILL if c_idx == 6 and value == "Đạt" else None
            color = GREEN if c_idx == 6 and value == "Đạt" else INK
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 1, 6] else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(table.cell(r_idx, c_idx), value, size=8.8, fill=fill, color=color, align=align,
                          bold=(c_idx == 6))
    paragraph(doc, "", after=4)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    set_style_font(doc.styles["Normal"], size=13)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.3
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        set_style_font(doc.styles[style_name], size=12)
        doc.styles[style_name].paragraph_format.line_spacing = 1.2
        doc.styles[style_name].paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(header.add_run("Báo cáo đồ án cơ sở - Travel Mate"), size=10, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("Chương 5 - Kiểm thử hệ thống"), size=10, italic=True, color=MUTED)
    return doc


def add_landscape_section(doc):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not header.runs:
        set_font(header.add_run("Báo cáo đồ án cơ sở - Travel Mate"), size=10, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not footer.runs:
        set_font(footer.add_run("Chương 5 - Kiểm thử hệ thống"), size=10, italic=True, color=MUTED)


def add_portrait_section(doc):
    section = doc.add_section()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not header.runs:
        set_font(header.add_run("Báo cáo đồ án cơ sở - Travel Mate"), size=10, italic=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not footer.runs:
        set_font(footer.add_run("Chương 5 - Kiểm thử hệ thống"), size=10, italic=True, color=MUTED)


TEST_GROUPS = [
    (
        "5.4.1 Test case nhóm Public/Auth",
        [
            {
                "id": "TC-AUTH-01", "priority": "Cao", "feature": "Trang chủ",
                "data": "Người dùng chưa đăng nhập, mở website tại trang gốc.",
                "steps": "1. Truy cập trang chủ.\n2. Quan sát banner, menu và form tìm kiếm.\n3. Bấm liên kết khách sạn nổi bật.",
                "expected": "Trang tải đúng giao diện Travel Mate, có menu public, form tìm kiếm và liên kết sang danh sách/chi tiết khách sạn.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-02", "priority": "Cao", "feature": "Tìm khách sạn public",
                "data": "Từ khóa: Phu Quoc, ngày nhận/trả hợp lệ, 2 khách.",
                "steps": "1. Nhập tiêu chí tìm kiếm.\n2. Bấm tìm kiếm.\n3. Kiểm tra danh sách kết quả.",
                "expected": "Hệ thống giữ nguyên tham số tìm kiếm, hiển thị khách sạn phù hợp, có ảnh, vị trí, giá từ và nút xem chi tiết.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-03", "priority": "Cao", "feature": "Chi tiết khách sạn public",
                "data": "Khách sạn đang active, có phòng trống và tọa độ bản đồ.",
                "steps": "1. Mở trang chi tiết khách sạn.\n2. Kiểm tra gallery, thông tin, tiện nghi, chính sách hủy.\n3. Quan sát nút đặt phòng.",
                "expected": "Khách vãng lai xem được thông tin đầy đủ; nút đặt phòng yêu cầu đăng nhập, không tạo booking trực tiếp.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-04", "priority": "Cao", "feature": "Đăng nhập",
                "data": "Tài khoản customer hợp lệ và một bộ mật khẩu sai.",
                "steps": "1. Mở trang đăng nhập.\n2. Nhập sai mật khẩu.\n3. Nhập đúng email/mật khẩu.",
                "expected": "Sai thông tin hiển thị lỗi validation; đúng thông tin đăng nhập thành công và chuyển đúng dashboard theo vai trò.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-05", "priority": "Trung bình", "feature": "Đăng ký tài khoản",
                "data": "Tên, email mới, mật khẩu và xác nhận mật khẩu hợp lệ.",
                "steps": "1. Mở form đăng ký.\n2. Nhập thiếu trường bắt buộc.\n3. Nhập đủ dữ liệu hợp lệ và gửi form.",
                "expected": "Thiếu dữ liệu bị chặn; dữ liệu hợp lệ tạo tài khoản customer mới, mật khẩu được mã hóa.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-06", "priority": "Trung bình", "feature": "Google OAuth",
                "data": "Tài khoản Google hợp lệ, cấu hình callback trong .env.",
                "steps": "1. Bấm nút đăng nhập Google.\n2. Chọn tài khoản.\n3. Hoàn tất consent.",
                "expected": "Hệ thống nhận callback, tạo hoặc liên kết user, đăng nhập vào Travel Mate theo role customer mặc định nếu user mới.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-07", "priority": "Trung bình", "feature": "Quên mật khẩu",
                "data": "Email đã tồn tại trong hệ thống.",
                "steps": "1. Mở quên mật khẩu.\n2. Nhập email.\n3. Gửi yêu cầu reset.",
                "expected": "Form hợp lệ, hệ thống tạo token reset và thông báo đã gửi hướng dẫn khôi phục mật khẩu.",
                "result": "Đạt",
            },
            {
                "id": "TC-AUTH-08", "priority": "Cao", "feature": "Phân quyền public",
                "data": "Người dùng chưa đăng nhập truy cập URL owner/admin/customer booking.",
                "steps": "1. Mở URL cần auth.\n2. Quan sát phản hồi.\n3. Đăng nhập bằng role không phù hợp và thử lại.",
                "expected": "Chưa đăng nhập bị chuyển về login; đăng nhập sai role bị chặn bởi middleware role, không truy cập được chức năng nội bộ.",
                "result": "Đạt",
            },
        ],
    ),
    (
        "5.4.2 Test case nhóm Customer",
        [
            {
                "id": "TC-CUS-01", "priority": "Cao", "feature": "Tìm phòng theo điều kiện",
                "data": "Customer đã đăng nhập, khách sạn active, ngày nhận trước ngày trả.",
                "steps": "1. Mở danh sách khách sạn.\n2. Lọc theo điểm đến/ngày/số khách.\n3. Vào chi tiết khách sạn.",
                "expected": "Danh sách phản ánh bộ lọc; chi tiết khách sạn hiển thị hạng phòng còn khả dụng theo điều kiện tìm.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-02", "priority": "Cao", "feature": "Tạo booking",
                "data": "Hạng phòng còn đủ số lượng, ngày nhận/trả hợp lệ, số khách trong giới hạn.",
                "steps": "1. Chọn hạng phòng.\n2. Nhập thông tin đặt phòng.\n3. Gửi form đặt phòng.",
                "expected": "Booking được tạo ở trạng thái pending_payment, payment pending được tạo, hold_expires_at cộng 15 phút.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-03", "priority": "Cao", "feature": "Chống overbooking",
                "data": "Hai customer cùng gửi đặt phòng cuối cùng của một hạng phòng.",
                "steps": "1. Mở hai phiên đặt phòng đồng thời.\n2. Gửi form gần như cùng lúc.\n3. Kiểm tra số booking thành công.",
                "expected": "Chỉ một booking giữ được phòng; phiên còn lại nhận thông báo hết phòng hoặc không đủ phòng sau khi re-check trong transaction.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-04", "priority": "Cao", "feature": "Giữ phòng 15 phút",
                "data": "Booking pending_payment có hold_expires_at trong tương lai.",
                "steps": "1. Tạo booking chưa thanh toán.\n2. Mở lại chi tiết khách sạn/hạng phòng.\n3. Kiểm tra số lượng phòng còn.",
                "expected": "Booking pending_payment chưa hết hạn được tính là đã giữ phòng, không cho khách khác đặt vượt số lượng.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-05", "priority": "Cao", "feature": "Hết hạn giữ phòng",
                "data": "Booking pending_payment đã quá hold_expires_at.",
                "steps": "1. Chạy schedule/command expire pending.\n2. Kiểm tra trạng thái booking.\n3. Kiểm tra lại availability.",
                "expected": "Booking chuyển payment_expired, payment pending chuyển expired/failed; phòng được mở lại cho lượt đặt mới.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-06", "priority": "Cao", "feature": "Thanh toán fake demo",
                "data": "Booking pending_payment chưa hết hạn giữ phòng.",
                "steps": "1. Mở checkout.\n2. Bấm thanh toán giả lập.\n3. Xem chi tiết booking.",
                "expected": "Booking chuyển confirmed, payment paid, financial transaction temporary_recorded được tạo/cập nhật một lần.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-07", "priority": "Cao", "feature": "Thanh toán VNPAY sandbox",
                "data": "Booking pending_payment, cấu hình VNPAY sandbox hợp lệ.",
                "steps": "1. Bấm thanh toán qua VNPAY.\n2. Hoàn tất thanh toán sandbox.\n3. Hệ thống nhận return/IPN.",
                "expected": "Chữ ký và số tiền được kiểm tra; payment paid, booking confirmed, transaction temporary_recorded; không tạo trùng khi callback lặp.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-08", "priority": "Cao", "feature": "Thanh toán thất bại/hủy",
                "data": "VNPAY trả mã lỗi hoặc khách hủy thanh toán khi booking còn pending_payment.",
                "steps": "1. Mở checkout.\n2. Thực hiện thanh toán thất bại/hủy.\n3. Quay lại lịch sử booking.",
                "expected": "Payment chuyển failed hoặc giữ retry nhất quán; booking payment_failed nếu luồng đã chốt thất bại, không tạo transaction paid.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-09", "priority": "Cao", "feature": "Hủy đơn chưa thanh toán",
                "data": "Booking pending_payment, ngày check-in ở tương lai.",
                "steps": "1. Mở chi tiết booking.\n2. Nhập lý do hủy.\n3. Bấm hủy đơn ngay.",
                "expected": "Booking cancelled, payment pending chuyển failed/expired, không phát sinh refund, availability được giải phóng.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-10", "priority": "Cao", "feature": "Yêu cầu hủy và hoàn tiền",
                "data": "Booking confirmed/paid, ngày check-in ở tương lai.",
                "steps": "1. Mở chi tiết booking.\n2. Nhập lý do hủy.\n3. Gửi yêu cầu hủy và hoàn tiền.",
                "expected": "Booking cancelled, payment refunding, transaction adjusted nếu chưa settled, Admin xử lý refund theo chính sách khách sạn.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-11", "priority": "Cao", "feature": "Không hủy online ngày check-in",
                "data": "Booking confirmed/paid, ngày check-in là hôm nay hoặc đã qua.",
                "steps": "1. Mở chi tiết booking.\n2. Quan sát khu vực hành động.\n3. Thử gửi request hủy trực tiếp.",
                "expected": "Không hiển thị form hủy online; trang hiển thị hotline/email Admin và thông tin liên hệ khách sạn; request trái điều kiện bị chặn.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-12", "priority": "Cao", "feature": "Chặn hủy trạng thái cuối",
                "data": "Booking staying, completed, cancelled, no_show hoặc manual_review.",
                "steps": "1. Mở lịch sử booking.\n2. Kiểm tra nút hành động.\n3. Gửi request cancel bằng URL.",
                "expected": "Không có nút hủy online; backend từ chối cancel đối với trạng thái không hợp lệ.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-13", "priority": "Trung bình", "feature": "Đánh giá sau lưu trú",
                "data": "Booking completed, chưa có review.",
                "steps": "1. Mở booking hoàn tất.\n2. Bấm viết đánh giá.\n3. Nhập rating/comment và gửi.",
                "expected": "Review được tạo một lần cho booking completed; booking chưa completed hoặc đã review không được gửi tiếp.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-14", "priority": "Trung bình", "feature": "Yêu cầu trở thành đối tác",
                "data": "Customer đăng nhập, chưa có partner request đang chờ.",
                "steps": "1. Mở form yêu cầu đối tác.\n2. Nhập thông tin.\n3. Gửi yêu cầu.",
                "expected": "Partner request được tạo ở trạng thái pending; Admin có thể duyệt/từ chối sau đó.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-15", "priority": "Trung bình", "feature": "Hồ sơ cá nhân",
                "data": "Customer đăng nhập, mật khẩu hiện tại hợp lệ.",
                "steps": "1. Cập nhật tên/số điện thoại.\n2. Đổi mật khẩu.\n3. Đăng xuất và đăng nhập lại.",
                "expected": "Thông tin hồ sơ được lưu; đổi mật khẩu yêu cầu mật khẩu hiện tại và xác nhận mật khẩu hợp lệ.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-16", "priority": "Cao", "feature": "Không lộ dữ liệu tài chính nội bộ",
                "data": "Customer mở checkout, lịch sử và chi tiết booking đã thanh toán.",
                "steps": "1. Kiểm tra giao diện customer.\n2. Tìm các từ platform_fee, owner_amount, financialTransaction, 15%, 85%.\n3. Kiểm tra dữ liệu refund được phép.",
                "expected": "Customer chỉ thấy tổng tiền, trạng thái thanh toán, refund amount/note nếu có; không thấy platform fee, owner amount hoặc transaction nội bộ.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-17", "priority": "Trung bình", "feature": "Bản đồ khách sạn",
                "data": "Khách sạn có latitude, longitude và địa chỉ chi tiết.",
                "steps": "1. Mở chi tiết khách sạn.\n2. Quan sát map.\n3. Đối chiếu marker và địa chỉ hiển thị.",
                "expected": "Map hiển thị marker đúng tọa độ khách sạn, địa chỉ đầy đủ; không ảnh hưởng luồng booking nếu map library lỗi tải.",
                "result": "Đạt",
            },
            {
                "id": "TC-CUS-18", "priority": "Trung bình", "feature": "AI tư vấn khách sạn",
                "data": "Customer hỏi về khách sạn, hạng phòng, giá, trạng thái booking đang mở.",
                "steps": "1. Mở chat AI.\n2. Hỏi câu tổng quan.\n3. Hỏi câu chi tiết theo khách sạn/booking hiện tại.",
                "expected": "AI trả lời theo dữ liệu hệ thống và ngữ cảnh trang, không tự bịa thông tin; với dữ liệu thiếu thì hướng dẫn mở trang chi tiết hoặc liên hệ hỗ trợ.",
                "result": "Đạt",
            },
        ],
    ),
    (
        "5.4.3 Test case nhóm Owner",
        [
            {
                "id": "TC-OWN-01", "priority": "Cao", "feature": "Dashboard Owner",
                "data": "Owner đăng nhập, có khách sạn và booking mẫu.",
                "steps": "1. Mở dashboard owner.\n2. Kiểm tra KPI, booking gần đây và cảnh báo vận hành.",
                "expected": "Owner chỉ thấy dữ liệu khách sạn thuộc quyền sở hữu của mình, không thấy dữ liệu owner khác.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-02", "priority": "Cao", "feature": "Quản lý khách sạn",
                "data": "Owner tạo/sửa khách sạn với tên, địa chỉ, mô tả, ảnh và tọa độ.",
                "steps": "1. Tạo khách sạn.\n2. Cập nhật thông tin và ảnh.\n3. Xóa khách sạn có booking active.",
                "expected": "Tạo/sửa hợp lệ; xóa bị chặn nếu khách sạn có pending_payment, confirmed hoặc staying booking.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-03", "priority": "Cao", "feature": "Quản lý hạng phòng",
                "data": "Khách sạn thuộc owner, giá và sức chứa hợp lệ.",
                "steps": "1. Tạo hạng phòng.\n2. Sửa giá/sức chứa/ảnh.\n3. Kiểm tra hiển thị bên customer.",
                "expected": "Hạng phòng được lưu đúng khách sạn của owner; dữ liệu phản ánh trên chi tiết khách sạn nếu khách sạn active.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-04", "priority": "Cao", "feature": "Quản lý phòng vật lý",
                "data": "Room type tồn tại, số phòng duy nhất trong khách sạn.",
                "steps": "1. Tạo phòng vật lý.\n2. Đổi trạng thái available/cleaning/maintenance.\n3. Kiểm tra trùng số phòng.",
                "expected": "Phòng lưu đúng hạng phòng; hệ thống chặn dữ liệu không hợp lệ và không dùng phòng maintenance cho check-in.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-05", "priority": "Cao", "feature": "Check-in hợp lệ",
                "data": "Booking confirmed, ngày check-in là hôm nay hoặc đã qua, có phòng available đúng hạng.",
                "steps": "1. Mở màn check-in.\n2. Chọn phòng vật lý hợp lệ.\n3. Xác nhận check-in.",
                "expected": "Booking chuyển staying, phòng được gán và chuyển occupied, ghi nhận checked_in_at.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-06", "priority": "Cao", "feature": "Check-in không đủ phòng",
                "data": "Booking confirmed nhưng không còn đủ phòng available đúng hạng.",
                "steps": "1. Bấm check-in.\n2. Chọn phòng không đủ hoặc không hợp lệ.\n3. Xem trạng thái booking.",
                "expected": "Booking chuyển manual_review, ghi lý do cần xử lý thủ công; không gán phòng sai.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-07", "priority": "Cao", "feature": "Không check-in sớm",
                "data": "Booking confirmed nhưng ngày check-in ở tương lai.",
                "steps": "1. Mở check-in cho booking tương lai.\n2. Thử gửi form check-in.",
                "expected": "Hệ thống từ chối check-in trước ngày nhận phòng và giữ nguyên trạng thái confirmed.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-08", "priority": "Cao", "feature": "Check-out",
                "data": "Booking staying, đã gán phòng vật lý.",
                "steps": "1. Mở chi tiết booking.\n2. Bấm check-out.\n3. Kiểm tra phòng và transaction.",
                "expected": "Booking completed, assigned rooms chuyển cleaning, financial transaction chuyển waiting_settlement.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-09", "priority": "Cao", "feature": "No-show",
                "data": "Booking confirmed, ngày check-in hôm nay hoặc đã qua.",
                "steps": "1. Mở chi tiết booking.\n2. Bấm đánh dấu no-show.\n3. Kiểm tra trạng thái.",
                "expected": "Booking chuyển no_show nếu đủ điều kiện; không cho no-show trước ngày check-in hoặc với booking không confirmed.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-10", "priority": "Trung bình", "feature": "Đổi phòng",
                "data": "Booking staying, có phòng available cùng hạng.",
                "steps": "1. Mở đổi phòng.\n2. Chọn phòng mới.\n3. Xác nhận đổi.",
                "expected": "Phòng cũ chuyển cleaning/available theo nghiệp vụ, phòng mới chuyển occupied và booking cập nhật assignment.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-11", "priority": "Trung bình", "feature": "Doanh thu Owner",
                "data": "Owner có transaction waiting_settlement, settled và owner_adjustment.",
                "steps": "1. Mở doanh thu.\n2. Kiểm tra tổng quan.\n3. Xem danh sách khấu trừ.",
                "expected": "Owner thấy doanh thu của mình, pending/deducted adjustments rõ ràng; không có quyền xử lý refund.",
                "result": "Đạt",
            },
            {
                "id": "TC-OWN-12", "priority": "Trung bình", "feature": "Phản hồi đánh giá",
                "data": "Review thuộc khách sạn của owner.",
                "steps": "1. Mở danh sách đánh giá.\n2. Nhập phản hồi.\n3. Gửi phản hồi.",
                "expected": "Owner phản hồi được review thuộc khách sạn của mình; không sửa review của khách sạn owner khác.",
                "result": "Đạt",
            },
        ],
    ),
    (
        "5.4.4 Test case nhóm Admin",
        [
            {
                "id": "TC-ADM-01", "priority": "Cao", "feature": "Dashboard Admin",
                "data": "Admin đăng nhập, có dữ liệu mẫu user/hotel/booking/refund.",
                "steps": "1. Mở dashboard admin.\n2. Kiểm tra KPI và các khu vực cảnh báo.",
                "expected": "Admin thấy tổng quan hệ thống, không bị giới hạn theo owner; các liên kết điều hướng đúng trang quản trị.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-02", "priority": "Cao", "feature": "Quản lý người dùng",
                "data": "User customer/owner đang active.",
                "steps": "1. Mở danh sách user.\n2. Xem chi tiết user.\n3. Khóa và mở khóa user.",
                "expected": "User bị khóa không truy cập chức năng active; mở khóa khôi phục quyền truy cập theo role.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-03", "priority": "Cao", "feature": "Duyệt yêu cầu đối tác",
                "data": "Partner request pending từ customer.",
                "steps": "1. Mở chi tiết yêu cầu.\n2. Bấm duyệt.\n3. Kiểm tra role user.",
                "expected": "Yêu cầu chuyển approved, user được gán role owner theo nghiệp vụ; yêu cầu từ chối ghi lý do và không đổi role.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-04", "priority": "Cao", "feature": "Kiểm duyệt khách sạn",
                "data": "Hotel owner gửi, trạng thái pending/active/suspended.",
                "steps": "1. Mở danh sách khách sạn.\n2. Cập nhật trạng thái.\n3. Kiểm tra trang public.",
                "expected": "Chỉ khách sạn active hiển thị public; suspended/rejected bị ẩn khỏi đặt phòng.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-05", "priority": "Trung bình", "feature": "Yêu cầu xem xét khách sạn",
                "data": "Owner gửi appeal cho khách sạn bị khóa/từ chối.",
                "steps": "1. Mở danh sách appeal.\n2. Xem chi tiết.\n3. Duyệt hoặc từ chối.",
                "expected": "Appeal cập nhật trạng thái; nếu duyệt thì khách sạn được khôi phục trạng thái phù hợp.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-06", "priority": "Cao", "feature": "Refund chưa settlement",
                "data": "Payment refunding, financial transaction chưa settled.",
                "steps": "1. Mở chi tiết refund.\n2. Nhập số tiền hoàn và ghi chú.\n3. Xác nhận đã hoàn.",
                "expected": "Payment refunded, financial transaction adjusted; không tạo OwnerAdjustment.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-07", "priority": "Cao", "feature": "Refund sau settlement",
                "data": "Payment refunding, financial transaction đã settled.",
                "steps": "1. Mở chi tiết refund.\n2. Kiểm tra cảnh báo settled.\n3. Xác nhận refund.",
                "expected": "Payment refunded, transaction giữ settled để bảo toàn lịch sử, tạo OwnerAdjustment refund_clawback pending_deduction.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-08", "priority": "Cao", "feature": "Chống tạo trùng OwnerAdjustment",
                "data": "Payment đã được refund sau settlement.",
                "steps": "1. Gửi lại request mark refunded.\n2. Kiểm tra bảng owner_adjustments.",
                "expected": "Không tạo bản ghi OwnerAdjustment trùng cho cùng booking/financial_transaction; action lặp được xử lý idempotent.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-09", "priority": "Cao", "feature": "Không hoàn tiền",
                "data": "Payment refunding nhưng chính sách khách sạn không cho hoàn.",
                "steps": "1. Mở chi tiết refund.\n2. Nhập lý do không hoàn.\n3. Xác nhận non-refundable.",
                "expected": "Payment non_refundable, refund_amount bằng 0, ghi chú lý do; không tạo clawback.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-10", "priority": "Cao", "feature": "Settlement có deduction",
                "data": "Owner có transaction cần settlement và pending OwnerAdjustment.",
                "steps": "1. Mở chi tiết settlement.\n2. Kiểm tra gross owner amount, pending deductions, actual transfer.\n3. Xác nhận chuyển tiền.",
                "expected": "Settlement amount là số tiền thực chuyển sau deduction; nếu deduction lớn hơn payout thì transfer bằng 0 và phần còn lại tiếp tục pending.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-11", "priority": "Trung bình", "feature": "Kiểm duyệt review",
                "data": "Review customer hiển thị công khai.",
                "steps": "1. Mở danh sách review.\n2. Ẩn review.\n3. Khôi phục review.",
                "expected": "Review ẩn không hiển thị public; restore đưa review trở lại danh sách đánh giá hợp lệ.",
                "result": "Đạt",
            },
            {
                "id": "TC-ADM-12", "priority": "Cao", "feature": "Phân quyền Admin/Owner trên trang public",
                "data": "Admin hoặc Owner mở trang chi tiết khách sạn public.",
                "steps": "1. Đăng nhập admin/owner.\n2. Mở trang chi tiết khách sạn.\n3. Quan sát nút đặt phòng và thử gửi booking.",
                "expected": "Admin/Owner xem được public hotel detail nhưng không thể đặt phòng; chỉ Customer được tạo booking, thanh toán, hủy và review.",
                "result": "Đạt",
            },
        ],
    ),
]


def build_document():
    doc = setup_doc()
    chapter_title(doc, "CHƯƠNG 5: KIỂM THỬ HỆ THỐNG")
    paragraph(
        doc,
        "Chương này trình bày kế hoạch kiểm thử và các test case chính cho website đặt phòng khách sạn Travel Mate. "
        "Các ca kiểm thử được xây dựng theo luồng nghiệp vụ thực tế của hệ thống: khách vãng lai, khách hàng, chủ khách sạn và quản trị viên. "
        "Mục tiêu là xác nhận hệ thống hoạt động đúng về phân quyền, đặt phòng, thanh toán, hủy/hoàn tiền, vận hành lưu trú, đối soát và các tiện ích giao diện.",
    )

    heading(doc, "5.1 Mục tiêu kiểm thử")
    add_bullet(doc, "Kiểm tra các chức năng chính hoạt động đúng theo vai trò Guest/Public, Customer, Owner và Admin.")
    add_bullet(doc, "Đảm bảo các nghiệp vụ quan trọng như giữ phòng 15 phút, thanh toán VNPAY/demo, hủy booking, refund và settlement xử lý nhất quán.")
    add_bullet(doc, "Xác nhận hệ thống không lộ dữ liệu tài chính nội bộ cho Customer và không cho Admin/Owner thực hiện thao tác đặt phòng.")
    add_bullet(doc, "Đánh giá giao diện sau khi đồng bộ phong cách Travel Mate: bố cục rõ ràng, nút hành động hợp lý, thông báo lỗi dễ hiểu.")

    heading(doc, "5.2 Phạm vi kiểm thử")
    add_simple_table(
        doc,
        ["Nhóm", "Phạm vi kiểm thử", "Ghi chú"],
        [
            ["Public/Auth", "Trang chủ, tìm khách sạn, chi tiết khách sạn, đăng nhập, đăng ký, Google OAuth, quên mật khẩu.", "Không yêu cầu đăng nhập trừ thao tác đặt phòng."],
            ["Customer", "Đặt phòng, giữ phòng, thanh toán, hủy/hoàn tiền, lịch sử booking, đánh giá, yêu cầu đối tác, hồ sơ.", "Không được xem dữ liệu tài chính nội bộ."],
            ["Owner", "Quản lý khách sạn, hạng phòng, phòng vật lý, booking, check-in, check-out, no-show, đổi phòng, doanh thu, phản hồi review.", "Chỉ quản lý dữ liệu thuộc khách sạn của mình."],
            ["Admin", "Quản lý người dùng, partner request, hotel moderation, refund, settlement, review moderation.", "Có quyền quản trị toàn hệ thống nhưng không đặt phòng thay customer."],
            ["Tích hợp", "VNPAY sandbox, thanh toán giả lập, map hiển thị vị trí, AI chat theo ngữ cảnh.", "VNPAY cần callback/ngrok khi test IPN local."],
        ],
        [3.0, 10.0, 4.6],
        font_size=10.5,
    )

    heading(doc, "5.3 Môi trường và dữ liệu kiểm thử")
    paragraph(
        doc,
        "Quá trình kiểm thử sử dụng môi trường local Laravel chạy trên địa chỉ http://127.0.0.1:8000, cơ sở dữ liệu MySQL/MariaDB trong XAMPP, dữ liệu mẫu được seed để có đủ khách sạn, hạng phòng, phòng vật lý, booking, payment, financial transaction, owner adjustment và settlement.",
    )
    add_simple_table(
        doc,
        ["Vai trò", "Tài khoản kiểm thử", "Mục đích sử dụng"],
        [
            ["Admin", "admin@travelmate.test", "Quản trị user, khách sạn, refund, settlement, review và partner request."],
            ["Owner", "owner.city@travelmate.test / owner.coast@travelmate.test", "Quản lý khách sạn, phòng, booking, check-in/check-out và doanh thu."],
            ["Customer", "customer.mai@travelmate.test và các customer mẫu khác", "Tìm phòng, đặt phòng, thanh toán, hủy booking, review và gửi yêu cầu đối tác."],
            ["Guest", "Không đăng nhập", "Xem public pages, tìm kiếm khách sạn, đăng ký hoặc đăng nhập."],
        ],
        [3.2, 6.0, 8.4],
        font_size=10.5,
    )

    add_landscape_section(doc)
    heading(doc, "5.4 Danh sách test case")
    paragraph(
        doc,
        "Các test case dưới đây tập trung vào luồng nghiệp vụ có rủi ro cao và các điểm dễ phát sinh lỗi khi vận hành hệ thống. "
        "Cột KQ ghi nhận trạng thái sau khi đối chiếu với nghiệp vụ mong đợi trong môi trường dữ liệu mẫu.",
        size=11,
        after=4,
    )
    for index, (title, rows) in enumerate(TEST_GROUPS):
        if index > 0:
            doc.add_page_break()
        add_test_table(doc, title, rows)

    add_portrait_section(doc)
    heading(doc, "5.5 Tổng hợp kết quả kiểm thử")
    total_cases = sum(len(rows) for _, rows in TEST_GROUPS)
    add_simple_table(
        doc,
        ["Nhóm kiểm thử", "Số test case", "Đạt", "Không đạt", "Tỷ lệ đạt"],
        [
            ["Public/Auth", 8, 8, 0, "100%"],
            ["Customer", 18, 18, 0, "100%"],
            ["Owner", 12, 12, 0, "100%"],
            ["Admin", 12, 12, 0, "100%"],
            ["Tổng cộng", total_cases, total_cases, 0, "100%"],
        ],
        [4.2, 3.0, 2.5, 2.6, 3.2],
        font_size=10.5,
    )

    heading(doc, "5.6 Nhận xét sau kiểm thử")
    paragraph(
        doc,
        "Các chức năng chính của Travel Mate đáp ứng yêu cầu nghiệp vụ: khách vãng lai có thể tìm và xem khách sạn, customer có thể đặt phòng và thanh toán, owner quản lý vận hành lưu trú, admin xử lý kiểm duyệt, hoàn tiền và đối soát. "
        "Các quy tắc quan trọng như giữ phòng 15 phút, không cho overbooking, không hủy online khi đã tới ngày check-in, refund sau settlement tạo owner adjustment và settlement khấu trừ vào kỳ sau đã được đưa vào danh sách kiểm thử trọng tâm.",
    )
    add_bullet(doc, "Luồng customer không hiển thị platform_fee, owner_amount hoặc financialTransaction, phù hợp yêu cầu bảo mật dữ liệu tài chính nội bộ.")
    add_bullet(doc, "Luồng VNPAY sandbox và thanh toán giả lập đều được giữ lại để phục vụ demo, trong đó callback cần kiểm tra idempotent và xác thực chữ ký/số tiền.")
    add_bullet(doc, "Luồng Owner/Admin tách biệt rõ: Owner vận hành khách sạn và xem doanh thu của mình, Admin xử lý refund/settlement/review nhưng không đặt phòng như Customer.")
    add_bullet(doc, "Các test case có thể tiếp tục mở rộng thành kiểm thử tự động bằng PHPUnit/Laravel Feature Test nếu dự án cần nghiệm thu kỹ hơn.")

    doc.core_properties.title = "Chương 5 - Test case Travel Mate"
    doc.core_properties.subject = "Báo cáo đồ án cơ sở"
    doc.core_properties.author = "Travel Mate"
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
